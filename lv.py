import streamlit as st
import google.generativeai as genai
import json
import os
from datetime import datetime
from docx import Document
import io
import tempfile
import re

# --- Utility to prettify keys ---
def prettify_key(key):
    """Converts a camelCase or snake_case key into a title-cased string."""
    key = key.replace('_', ' ')
    key = re.sub(r'([a-z])([A-Z])', r'\1 \2', key)
    return key.title() + ":"

# --- Lee Valley Golf Club Minutes Generator ---
def generate_golf_club_minutes(structured):
    """Generates meeting minutes text based on the Lee Valley Golf Club template."""
    now = datetime.now()
    # Helper to get value or fallback
    def get(val, default="Not mentioned"):
        return val if val and val != "Not mentioned" else default

    # Helper for formatting bullet points or lists of items
    def format_items(val):
        if isinstance(val, list) and val:
            return "".join([f"• {item}\n" for item in val])
        elif isinstance(val, str) and val.strip() and val != "Not mentioned":
            return f"• {val}\n"
        else:
            return "• Not mentioned\n"

    # --- Extract data from the structured JSON ---
    title = get(structured.get("titleOfMeeting"), "Lee Valley Mens Club Committee Meeting")
    purpose = get(structured.get("purposeOfMeeting"))
    location = get(structured.get("locationOfMeeting"), "Lee Valley")
    attendees = structured.get("attendees", [])
    apologies = structured.get("apologies", [])
    meeting_date_time = get(structured.get("meetingDateTime"), now.strftime("%d/%m/%Y @ %H:%M"))
    next_meeting_date_time = get(structured.get("nextMeetingDateTime"))
    prepared_by = get(structured.get("minutesPreparedBy"))
    date_circulated = get(structured.get("dateCirculated"), now.strftime("%d/%m/%Y"))
    circulation = get(structured.get("circulation"))

    # --- Meeting body items ---
    training = format_items(structured.get("training", []))
    health_safety = format_items(structured.get("healthAndSafety", []))
    finance = format_items(structured.get("finance", []))
    issues_risk_discipline = format_items(structured.get("issuesRiskDiscipline", []))
    teams = format_items(structured.get("teams", []))
    projects = format_items(structured.get("projects", []))
    competitions = format_items(structured.get("competitions", []))
    comments = format_items(structured.get("comments", []))
    aob = format_items(structured.get("anyOtherBusiness", []))
    captains_comments = format_items(structured.get("captainsClosingComments", []))

    # --- Compose the minutes string ---
    template = f"""
Title of Meeting: {title}
Purpose of Meeting: {purpose}
Location of Meeting: {location}
Date / Time of Meeting: {meeting_date_time}
Date / Time of Next Meeting: {next_meeting_date_time}
Minutes Prepared By: {prepared_by}
Date Circulated: {date_circulated}
Circulation: {circulation}

________________________________________
ATTENDEES:
{format_items(attendees)}
________________________________________
APOLOGIES:
{format_items(apologies)}
________________________________________

MEETING MINUTES & ACTIONS
________________________________________

1. Training (First Aid, Programmes, etc.)
{training}
________________________________________
2. Health and Safety
{health_safety}
________________________________________
3. Finance (Status, Projections)
{finance}
________________________________________
4. Issues, Risk, Discipline
{issues_risk_discipline}
________________________________________
5. Teams (Purcell, Bruen, etc.)
{teams}
________________________________________
6. Projects (Defib, Simulator, 5 Year Vision)
{projects}
________________________________________
7. Competitions (Weekly, Matchplays)
{competitions}
________________________________________
8. Comments
{comments}
________________________________________
9. Any Other Business (AOB)
{aob}
________________________________________
10. Captain's Closing Comments
{captains_comments}
"""
    return template.strip()

# --- DOCX Export Functions (Updated for Lee Valley) ---
def create_narrative_docx(narrative_text):
    """Creates a DOCX for the narrative summary."""
    doc = Document()
    doc.add_heading("Lee Valley Golf Club - Meeting Summary", level=1)
    doc.add_paragraph(narrative_text)
    output = io.BytesIO()
    doc.save(output)
    output.seek(0)
    return output

def create_keypoints_docx(text):
    """Creates a DOCX for key points and actions."""
    doc = Document()
    doc.add_heading("Lee Valley Golf Club - Key Points & Actions", level=1)
    doc.add_paragraph(text)
    output = io.BytesIO()
    doc.save(output)
    output.seek(0)
    return output

def create_minutes_docx(content):
    """Creates a formatted DOCX for the final minutes."""
    doc = Document()
    doc.add_heading("Lee Valley Golf Club Meeting Minutes", level=1)
    # Simple line-by-line processing, can be enhanced for better formatting
    for line in content.splitlines():
        if line.strip().endswith(":") and not line.startswith("•"):
             # Simple check for headings
            try:
                doc.add_heading(line.strip(), level=2)
            except Exception:
                doc.add_paragraph(line) # Fallback for any heading issues
        elif line.strip() == "________________________________________":
            doc.add_paragraph("--------------------------------------------------")
        elif line.strip():
            doc.add_paragraph(line)

    output = io.BytesIO()
    doc.save(output)
    output.seek(0)
    return output

# --- Configure Gemini API ---
try:
    # It's recommended to use st.secrets for API keys
    genai.configure(api_key=st.secrets["GEMINI_API_KEY"])
    model = genai.GenerativeModel(model_name='gemini-1.5-flash') # Updated model
except KeyError:
    st.error("GEMINI_API_KEY not found in Streamlit secrets. Please add it to continue.")
    st.stop()
except Exception as e:
    st.error(f"Error configuring Gemini API: {e}")
    st.stop()

st.set_page_config(page_title="LVGC Minutes", layout="wide")

# --- Logo Data ---
logo_base64 = "data:image/jpeg;base64,/9j/4AAQSkZJRgABAQAAAQABAAD/2wCEAAkGBxEQEBUQEBIVFRUWEBcQGRcYDw8VFhEVHxIWFhUdHhUYHSggGBolHxMZLTEhJTUtLi4uFx8zODMsNygtLisBCgoKDQ0OFRANFSsdFx0rKy0tKy03LSstNystLS0rLSs3Nys3LSstKystLS0rNy0rKystKysrKysrKy0tKy0rK//AABEIAKAAnQMBIgACEQEDEQH/xAAbAAACAwEBAQAAAAAAAAAAAAAABQEEBgcDAv/EAEkQAAIBAwICBQgFBwkJAQAAAAECAwAEERIhBTEGEyJBUQcUFmFxgZTSIzKRodE0UlNydLPBFTNCVGJjk7GyFyU1Q5KitMLwJP/EABgBAQADAQAAAAAAAAAAAAAAAAABAgME/8QAIREBAQADAAIDAAMBAAAAAAAAAAECESExUQMSQRNhgQT/2gAMAwEAAhEDEQA/AO40UUUBRRRQFFFFAUUUUEUu4/xE2tu84XVp07atOcuF5++mOaRdMo9VlKNQXOgZPLPWJjPh4VXO2S6TPPWc/wBorf1Yf4rfLR/tFb+rD/Fb5axM0TIxVhgiviuP+XOfro+mLoHD+nbSzRxdQBrkRM9bnTqIGcafXW5Fcg6PWJE8Ekh0AzppyMtIdYxgeGebcvadq6+K3+HLLKW5Mc5JeJooordQUUUUBRRRQFFFRQTRRRQFFFFAVU4ncdVDJKN9EbP7SFJxVulXH49caxd0kyI2+MrnWR7wmPfUZeCMRf3M9yeohWV5lkDu/W9gEZA0jOEG/t276i8jl8ymSeWcSxhGaJ2RkdWmGCDgkj38xTHhPFIrC4uLefshpjKrAM2VPIHG/LH3148avDxGUw2o283bLOrLrHWI4x4boME+Jrmslm999NZ5/pk7e6UqI5hlO4j68fs8R/ZP3c6tvbx2yLJtMXyUOluqTHiD9d/7J2HfmlUiFSVYYIJUjvBHMVGs4xnbOcZ2z41h9vbbW1/h08kl3E+rVIZ0wWLY1axjOO72V08z3684YX/VmdT961y/gH5Xb/tEf7wV2quj4JuVj8vLGXv+lLWxUXNs6Bs4KypIMDGeWPzhWmRsgH31XveHxTgCVFcKdQyM4NWQK3ksvazun1RRRVkCioqaAooooCiiigKKKKApfxu1MsLKG0MMOrfmsp1A/dTCvhlBGDy5e2os3BgOjfB4p4klhmKzqzGRtOskt3EH1cj7a1XAeDi1VgX6x2cuXIwzZ7j9lZrpZ0e83j6+z1R4OZAruAVzs3Pu/wAvZXt0Zke8ZZZbxiY8Ewr2ACPzvzx/96qxx5da60vZvfGS6VppvZh/eavtAP8AGlNa3yk24W5RwPrx7+sg4/yxWSrm+Saysb4XeMX+Afldv+0R/vBXaq4rwD8rt/2iP94K7VXR/wA/isfl8xNFRU10MhRRRQFFRU0BRRRQRU0UUBUVNFAUUUUFTiZAhkJ5dWxPrGk1yXghjUO5maOVQGi0q7FmHcQBjB2G5ro/TO8ENnIe9x1Q9ZbY/dk+6s95Mnj+lX/mdlvanq9539ornz7nI0x5jap9JWu7m2jkltdAjGovqXJyAP5vmg76yRQgZ/jXZePyhbWYldX0TDGPFSKi64FbyxrGyDSgwuOyVHqIqM/h+18px+TU8OVdH0PnUB7vOI+/n9IK7RS+z4TDCulEAGoN4ksDkEk77UwrX48PrFc8vtU0UUVooKKKKAoqKmgiipooCoqaKAooooCoNTXlNKFUsxwACx9QHOg5p5QeKmWfqB9WL/uc8z7ht9tKOjd8YLqNxy1BW9anY/j7qfDhYvLW5ul3c3DyIcdrQv8AR+wnb1CvPhVmh4TPLjtdYN8bjSUIx/1H7a47jlc9/wCuiWTHTe8ZUNA4PJgF+1gP41fFJLjiEc9m0yHUuA/P6pUhiD4YxTlHBGQcjnz511Sy1zvuoqaKsCiiigKipooCiiigiipooCiiigKKKKCKVdKQfMp9PPqm+zHa+7NNao8cA82m1cuofPs0Gq5eKmeWWs2kis0FtbsslwAu0i6Vbq9pBk94Gfdv62lvwExcPkt1OXeNyTnZpCv+WwFeHRktJ5uTyitAR62digPuWM/bWqquGMs2m3rnPBeMLZlYuqcBlHWxlH1o4GC4zsUPhWv6NvqiLICIy7GMHmI+72DOcDwxVXpPG0Wi8jBLRHDgf04T9ce7n6t6tcAmGl0UgqJC6euOT6QH2Zcj3VGMuN1S9mziiiitVRRRRQFRU1FBNFRU0C+/4tFAQsmrcZ2ilf8A0A1V9Jrbxk+GuPlqt0q6Sfyehmkt5ZIVALSIYToJbABV2B5439Yqz0d4y15EJuoeKNlV4y7REyKRnOEJ093PxqtmSeD0mtvGT4a4+Wj0mtvGT4a4+WnBNBp32cJ/Sa28ZPhrj5aPSa28ZPhrj5ab5qSad9nCdekluTgGTnj8nnH/AK0h6W9I42DWkTjJyrth8KBzUYBJJ5csVrb2V0jZ0jMjAZCBkUt6gWwB76y/RvprHfzvDBaTAxtplL+bp1J35jrCTuCNhVcscrNbTLJdrfRu6MkvYhZIo4FiDHUBJgjRgEDl2vtrUUg6U9IP5Pi694JJIgMuyNCOr3AHZdgTknuq/wAIvWnhWVoni1DUEfRrA/o50kgEjBx3VbGaiKuSIGBBGQQRjxrGG0ntSqRxM3VglJFVm6wGQnq3A5Jg8zyO4rbZoNMsdkuiq542kblGjmJGN1t5XU7Z2IG9efpFH+iufhJfwpzUZpq+zhP6RR/orn4SX8KPSKP9Fc/CS/hTgGgGmr7OE/pFH+iufhJfwo9Io/0Vz8JL+FOCRQCKavs4Uw8dRmCiKcZIXJtpVUZ8SRsKb1GoV9VMl/UMd5Xf+DXXsi/8iOnHQz/h1n+xQfuUpF5YJ1XhE6sQCxjVRndj18ZwPHYE+6vXor0osI+HWoe7gUraQqwM8QZWESAgjOc5qRd6fWMcvD7hnXLR28syHODHIsLkMD3GsXwXhvDpOERNcuDNLEqk+cO7iSR9EZ6sSb4LqcU64r0oS64dxRkZGijjktkYN/OFrYZOeR7UmBjwpf0fj4PJwy1E0tpHKscMpbrLVJlkQh92O+SVwe/BNB5eVPhiW9jZsNpI5oLQyKzKTEI3yux5ZGacXHAuGSXFstuy9cky3C6ZpJAViILg9ogZyNzSjyjcZiueHWEsgCLLfQzaWOfogH1E+rBH204vr7hNu6Xdu9sZkBiWOCS1DTdY6JghNzjmPDeg3IFcrvv919IklG0F+ug9wEpIH+rSc/3hrqgrEeVzgxueHNIg+ktyLhSOYC/XA92T7UFB69MB53d2nDhuuvz2fw6mM9hT6mcge6tjnArD+TEy3UcnE7gDrLgpGo7liiGjb2trP2Vc8oPGVit/NkkVJrlltlyygxq5w7nfYBdW/jigykfHpk4rDxF2Pml272CeCIjgRue4an1H2Zrd9MrJJ7GcOPqQySoc4KSCJ8MD4jJrKdKuhzHhjR+fu0cMQeJXjslQMg+jHWJGCM8s57++rPDOlcd3wKWdnUyJZSJKupdQcRsMn9bYj9agYeS+0ReGwSgduWMM7EsWch3xknwyaUeWi2VbIXCgrL18cWtWZW0dvs7HluaueT3pDZR8Ltke6gRli0lWniVlIc7EE5pH5UOOR3XB45lwBJe6U7WrrFR5F1D1EJn1ZFB03h9lHBGI4lCIMnA5ZJyfvJrm/Szh9uOPWiSYSOaKWSb6Vo1kYJJoLEEb5ArXcU6Z2UEEkouYJGRCwRbmItIwGygAncnasp0wmgPHuH+dGIRi2lLiQx6FJjkxq17c8Y9dA4sej9hLemS2IPV27xShZHcOJvqYfUdBAR+W/bpPYBuCcU6hmPmd6SULMxEE47iT48s9+R4U+4ffcPgu0SxaBnu3CMkMsGlBHDNIX0R9/d68jwpZ5bZI/wCTGVmXX1sbKpI1fWwSBz5E0DDo/Ab+8PE5M9VGGhtFzgFeUkpH9ojA9Q9lbSqHBpImgj6kqUEagaCpUAKNhir9BVvIInAMqowU6u0qsE2578ts70utbjh8raImtnYgsFVoGYgd+Bvil/lDidrRdKs8a3UMk6KGLPbrIDKNI3O3MeApL0r40l0nm3DxFM7WkxR45syQYj3XqwOwWB0jJG+fCg1NrNw5ibeJrUksWMSmAksp5mMd4I547q+7VrCZWMXm0irsxXqGVf1scu/nWU43eWt3ZwQWWnrxLA0Ua/zloVkTWWHOIKuoEnHvzVzoxxu1t47jrJUTVxScAalDHXOEDYz9TLZzyxvQPLW94dMyRRSWsjKpCIr27Mq94VRyHZ7vCrEtvaK6I6Qh2OUUrEGcrv2QeZHPblWU8n0MrgOJIjCk132BHiRJWucoxfJyCpY7BdnHPmEPS7jLRcSe47Gq2HVJCUZn2jZ459GV6xCZZFOD2dmOcUHXCcUsh4vaTMYUuIZGIYaBNEzMB9bsA5Pfms7x2aa84LI0MizuyKSYFdFkAkQyKo1Enshl57+qvLj/ABG1u7B4rF4zOIH6mNezLGRGQyiMdqM6Cy7450GmsuKWeRBDNBlQQI0liyoHMBAdsV4teWEqvOZLZ1UBXkLQMEGdgz93M7HxrLcbvLW8tIIbHT16ywNFGuBLa6ZE6wuOcQC6gc49+a8bPikMNrxC0kYC4e6vNEOO3L1jOY9Kc3ByNxtQbiG5tbgGJGhlVQpZFaJwo5plRnHLI9lL34nwsZDTWYBODmS1AYqc4O++Dv6jWf4pZsb+4vLIg3MAjHV6triHq/pIyO45Awe4gUv4pNFNwO5uVXS000yrqXTJg3rvoxzyN9vUaDYHi3CcZ6+yxy/nrXf76s27WLlIY/N2KxmREXqSVjY7sqjkpONxtWS6YuRe2j2DQGZo7qUatLI7PDEsZOCN2CYU8jjvxVO6S3j4Vb3Vq5intpdKGRQHWR3/AP0Rum2Rh2OkdyAjYUS3y2dpKGAjhcKxRuxEwVh3HbYjbaqsnEuGs5LTWhfByTJblsLzzvnb7qX8d4WV4PLBYkuxg1KVbLzZIZ21j67uNRz3lq+7PjPDJIo1R4h1alliPZeDEZB1R80wpI3FEGccliEF0pt9C5xMDDpXOxxINh4c6Li9sWEcskluQ4KxuzwnrBnkjH62/cKzHC+JxHg0Vorhpn4aYxGp1OGFqSQQNwdsb95r14fxq2uUsIbeRXeN0d1HOFUtpA+sf0N8DfvIoNJwvidk56q1mt2O7aIpYifWcIfvptXJ+jAzc2DRuk2ma7zGsYzbI5kPWNID37ABufWbeNdXoKl9fJCutw2M47KOxz7APVVD0kt/CX4af8KdYoxVbL+J4Teklv4S/DT/AC0ekkHhL8NP8tOcUYpq+zhL6S2/hL8NP8tHpJb+Evw0/wAtOsUYpq+zhKOkkHhL8NP8tHpJb+Evw0/y06xRimr7OEvpJb+Evw0/y0eklv4S/Cz/AC06xRimr7OE3pLB4S/DT/LR6SW/hL8NP8tOcUYpq+zhN6SweEvw0/y0eklv4S/Dz/hTnFGKavs4TekkHhL8NP8ALUeklv4S/Cz/AC06xRimr7OE3pLB4S/DT/LUeklv4S/Cz/LTrFGKavs4TJ0igJAAl3OPyecfwpyDRijFTN/qH//Z"

# --- Password protection ---
if "password_verified" not in st.session_state:
    st.session_state.password_verified = False

if not st.session_state.password_verified:
    st.title("🔒 LVGC Minutes Recap Access")
    st.warning("This application requires a password to proceed.")
    with st.form("password_form"):
        user_password = st.text_input("Enter password:", type="password", key="password_input")
        submit_button = st.form_submit_button("Submit")
        if submit_button:
            try:
                expected_password = st.secrets["password"]
                if user_password == expected_password:
                    st.session_state.password_verified = True
                    st.rerun()
                else:
                    st.error("Incorrect password. Please try again.")
            except KeyError:
                st.error("Password not configured in Streamlit secrets. Please contact the administrator.")
            except Exception as e:
                st.error(f"An error occurred during password verification: {e}")
    st.stop()

# --- Sidebar ---
with st.sidebar:
    st.image(logo_base64, use_container_width=True)
    st.title("📒 LVGC Minutes Recap")
    
    if st.button("🔄 Restart Session"):
        keys_to_clear = ['transcript', 'structured', 'minutes', 'narrative', 'keypoints_summary']
        for key in keys_to_clear:
            if key in st.session_state:
                del st.session_state[key]
        st.rerun()

    if st.button("About this App", key="about_button_sidebar"):
        st.sidebar.info(
            "**LVGC Minutes Recap** helps generate meeting minutes for the Lee Valley Golf Club committee. "
            "Upload or record audio, and the app will transcribe and summarise it."
        )
    if st.button("Created by Dave Maher", key="creator_button_sidebar"):
        st.sidebar.write("This application's intellectual property belongs to Dave Maher.")
    st.markdown("---")
    st.markdown("Version: 2.0.2 (LVGC)")

# --- Main UI Header ---
col1, col2 = st.columns([1, 6])
with col1:
    st.image(logo_base64, width=80)
with col2:
    st.title("📝 LVGC Minutes Recap")
    st.markdown("#### Lee Valley Golf Club Minute-AI (MAI) Generator")

st.markdown("### 📤 Record or Upload Meeting Audio")

# --- Input Method Selection ---
mode = st.radio(
    "Choose input method:",
    ["Record using microphone", "Upload audio file"],
    horizontal=True,
    key="input_mode_radio"
)

audio_bytes = None

if mode == "Upload audio file":
    uploaded_audio = st.file_uploader(
        "Upload an audio file (WAV, MP3, M4A, OGG, FLAC)",
        type=["wav", "mp3", "m4a", "ogg", "flac"],
        key="audio_uploader"
    )
    if uploaded_audio:
        st.audio(uploaded_audio)
        audio_bytes = uploaded_audio

elif mode == "Record using microphone":
    # Note: Streamlit's audio_recorder is deprecated. Using a more generic term.
    # For a real app, consider a custom component if more control is needed.
    st.info("Recording functionality may vary by browser. Please use the upload feature for best results.")
    recorded_audio = st.audio_input("🎙️ Click the microphone to record, then click again to stop and process.", key="audio_recorder_main")
    if recorded_audio:
        st.audio(recorded_audio, format="audio/wav")
        audio_bytes = recorded_audio


# --- Transcription and Analysis ---
if audio_bytes and st.button("🧠 Transcribe & Analyse", key="transcribe_button"):
    with st.spinner("Processing with Gemini... This may take a few minutes for longer audio."):
        # Simplified audio data handling
        if hasattr(audio_bytes, "read"):
            audio_data_bytes = audio_bytes.read()
        else:
            st.error("Could not read audio data. Please try again.")
            st.stop()
        
        # Determine file extension
        file_extension = ".wav" # Default
        if hasattr(audio_bytes, 'name') and isinstance(audio_bytes.name, str):
            original_extension = os.path.splitext(audio_bytes.name)[1].lower()
            if original_extension:
                file_extension = original_extension

        with tempfile.NamedTemporaryFile(delete=False, suffix=file_extension) as tmp_file:
            tmp_file.write(audio_data_bytes)
            tmp_file_path = tmp_file.name
        
        try:
            st.info(f"Uploading audio to Gemini for processing...")
            audio_file_display_name = f"LVGC_Recap_{datetime.now().strftime('%Y%m%d_%H%M%S')}"
            audio_file = genai.upload_file(path=tmp_file_path, display_name=audio_file_display_name)
            st.success(f"Audio uploaded successfully: {audio_file.name}")

            prompt = (
                "You are an expert transcriptionist for Lee Valley Golf Club committee meetings. "
                "Transcribe in UK English the following meeting audio accurately. "
                "For each speaker, if a name is mentioned, use their name (e.g., Captain:, John Smith:). "
                "If not, label generically as Speaker 1:, Speaker 2:, etc., incrementing for each new unidentified voice."
            )
            result = model.generate_content([prompt, audio_file], request_options={"timeout": 1200})
            st.session_state["transcript"] = result.text
            st.success("Transcript generated successfully.")

        except Exception as e:
            st.error(f"An error occurred during transcription: {e}")
        finally:
            if 'audio_file' in locals() and audio_file:
                try:
                    genai.delete_file(audio_file.name)
                    st.info(f"Cleaned up uploaded file: {audio_file.name}")
                except Exception as del_e:
                    st.warning(f"Could not delete uploaded file {audio_file.name} from Gemini: {del_e}")
            if os.path.exists(tmp_file_path):
                os.remove(tmp_file_path)

# --- Display Transcript and Generate Minutes ---
if "transcript" in st.session_state:
    st.markdown("---")
    st.markdown("## 📄 Transcript")
    st.text_area("Full Meeting Transcript:", st.session_state["transcript"], height=300, key="transcript_display_area")

    if st.button("📊 Extract & Format Meeting Minutes", key="summarise_button"):
        with st.spinner("Generating structured meeting minutes..."):
            try:
                current_transcript = st.session_state['transcript']
                prompt_structured = f"""
You are an AI assistant for Lee Valley Golf Club meetings.
Your task is to extract detailed, structured information from the provided meeting transcript and return a JSON object.
Use UK English. Format all dates as DD/MM/YYYY and all times as HH:MM (24 hour).
If a key is not mentioned, use "Not mentioned" or an empty list [].

Keys to extract:
- titleOfMeeting
- purposeOfMeeting
- locationOfMeeting
- meetingDateTime
- attendees (list of names)
- apologies (list of names)
- minutesPreparedBy
- dateCirculated
- circulation (string describing who gets the minutes)
- nextMeetingDateTime
- training (list of key points/actions)
- healthAndSafety (list of key points/actions)
- finance (list of key points/actions)
- issuesRiskDiscipline (list of key points/actions)
- teams (list of key points/actions, e.g., Purcell, Bruen)
- projects (list of key points/actions, e.g., Defib, Simulator, 5 Year Vision)
- competitions (list of key points/actions, e.g., weekly, matchplays)
- comments (list of general comments made)
- anyOtherBusiness (list of AOB points)
- captainsClosingComments (list of points)

Transcript:
---
{current_transcript}
---

Provide ONLY the JSON object in your response. Do not include any other text.
"""
                response = model.generate_content(prompt_structured, request_options={"timeout": 600})
                
                # Clean and parse JSON
                json_text_match = re.search(r"```json\s*([\s\S]*?)\s*```|({[\s\S]*})", response.text, re.DOTALL)
                if json_text_match:
                    json_str = json_text_match.group(1) or json_text_match.group(2)
                    try:
                        structured = json.loads(json_str.strip())
                        st.session_state["structured"] = structured
                        # Generate minutes immediately after successful extraction
                        minutes_text = generate_golf_club_minutes(structured)
                        st.session_state["minutes"] = minutes_text
                        st.success("Meeting minutes generated in Lee Valley Golf Club format.")
                    except json.JSONDecodeError as e:
                        st.error(f"❌ Failed to parse JSON from AI response. Error: {e}")
                        st.code(json_str.strip(), language="json")
                else:
                    st.error("❌ No valid JSON object found in Gemini's response for structured summary.")
                    st.code(response.text)

            except Exception as e:
                st.error(f"An error occurred during summarization: {e}")

# --- Display Formatted Minutes and Download ---
if "minutes" in st.session_state:
    st.markdown("---")
    st.markdown("## ⛳ Lee Valley Golf Club Meeting Minutes (Draft)")
    st.text_area(
        "Drafted Meeting Minutes:",
        st.session_state["minutes"],
        height=900,
        key="minutes_text_area"
    )
    st.download_button(
        label="📥 Download Minutes (DOCX)",
        data=create_minutes_docx(st.session_state["minutes"]),
        file_name=f"LeeValleyGC_Minutes_{datetime.now().strftime('%Y%m%d_%H%M')}.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        key="download_minutes_docx"
    )

# --- Generate, Display, and Download Summaries ---
if "transcript" in st.session_state:
    st.markdown("---")
    st.markdown("## 🔍 Meeting Summaries")
    
    col1, col2 = st.columns(2)

    with col1:
        if st.button("📝 Generate Narrative Summary", key="narrative_button"):
            with st.spinner("Creating a narrative summary..."):
                try:
                    prompt_narrative = f"""
You are an AI assistant creating a professional, concise summary of a Lee Valley Golf Club committee meeting in UK English.
Based on the following transcript, write a coherent, narrative summary. The summary should be well-organized and capture the main points, discussions, and outcomes.

Transcript:
---
{st.session_state['transcript']}
---
Narrative Summary:"""
                    response = model.generate_content(prompt_narrative, request_options={"timeout": 600})
                    st.session_state["narrative"] = response.text
                except Exception as e:
                    st.error(f"Error generating narrative summary: {e}")

    with col2:
        if st.button("🧾 Generate Key Points & Actions", key="keypoints_button"):
            with st.spinner("Summarising for key points and actions..."):
                try:
                    prompt_keypoints = f"""
You are an AI assistant for Lee Valley Golf Club committee meetings.
Summarise the following transcript into concise bullet points, focusing on:
- Key discussion points
- Major decisions made
- All action items (with responsible persons/roles and deadlines, if mentioned)

Be succinct, avoid repetition, and use bullet points.

Transcript:
---
{st.session_state['transcript']}
---"""
                    response = model.generate_content(prompt_keypoints, request_options={"timeout": 600})
                    st.session_state["keypoints_summary"] = response.text
                except Exception as e:
                    st.error(f"Error generating key points summary: {e}")

    if "narrative" in st.session_state:
        st.markdown("### Narrative Summary")
        st.text_area("Meeting Narrative:", st.session_state["narrative"], height=400, key="narrative_text_area")
        st.download_button(
            label="📥 Download Narrative (DOCX)",
            data=create_narrative_docx(st.session_state["narrative"]),
            file_name=f"LVGC_Narrative_Summary_{datetime.now().strftime('%Y%m%d_%H%M')}.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            key="download_narrative_docx"
        )

    if "keypoints_summary" in st.session_state:
        st.markdown("### Key Points & Actions")
        st.text_area("Key Points & Actions:", st.session_state["keypoints_summary"], height=400, key="keypoints_text_area")
        st.download_button(
            label="📥 Download Key Points (DOCX)",
            data=create_keypoints_docx(st.session_state["keypoints_summary"]),
            file_name=f"LVGC_KeyPoints_{datetime.now().strftime('%Y%m%d_%H%M')}.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            key="download_keypoints_docx"
        )

# --- Footer ---
st.markdown("---")
st.markdown(
    "**Disclaimer:** This implementation is a draft. "
    "Adjustments may be required for optimal performance. "
    "Always verify the accuracy of AI-generated transcriptions and minutes."
)
st.markdown("Created by Dave Maher | For Lee Valley Golf Club internal use.")

